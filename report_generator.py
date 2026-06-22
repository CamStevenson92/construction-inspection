"""
Site inspection report generator — Meinhardt ANZ.

Auto-replaced template placeholders:
    <<SITE_NAME>>            Site / project name
    <<PROJECT_NUMBER>>       Project number
    <<INSPECTOR_NAME>>       Inspector full name
    <<INSPECTION_DATE>>      Earliest photo date
    <<REPORT_DATE>>          Today's date
    <<SITE_ADDRESS>>         Site address
    <<TOTAL_PHOTOS>>         Total photo count
    <<INSPECTION_WEATHER>>   Weather summary from photo EXIF / Open-Meteo API
    <<WORKS_SUMMARY>>        AI-generated ≤100-word summary of inspected works

When the Meinhardt ANZ .dotx template is provided:
  - All placeholders above are replaced automatically
  - AI-generated works summary is inserted after the "proposed works" paragraph
  - Site Observations table is populated from photo issues (one row per photo
    that has issues_found text, referencing the photo number)
  - Site Photographs appendix is replaced with actual photos in a 2×2 grid
    (4 photos per page); each caption shows metadata + three inspection fields
  - Yellow-highlighted user-input fields remain for manual completion
"""

import io
import os
import re
import shutil
import tempfile
import zipfile
from collections import Counter
from datetime import date
from typing import List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage
from PIL import ImageOps

from models import PhotoData

try:
    import anthropic as _anthropic
    _ANTHROPIC_AVAILABLE = True
except ImportError:
    _ANTHROPIC_AVAILABLE = False

PLACEHOLDER_RE = re.compile(r"<<([A-Z_]+)>>")

# Maximum photo dimensions (inches) inside the 2-col photo grid cell
PHOTO_MAX_W_IN = 3.3
PHOTO_MAX_H_IN = 2.9

# Caption colours matching the Meinhardt template
COL_NAV       = RGBColor(0x02, 0x3F, 0x88)  # Meinhardt navy (header blue)
COL_INSPECTED = RGBColor(0x00, 0x4F, 0x9F)  # blue
COL_ISSUES    = RGBColor(0xCC, 0x44, 0x00)  # burnt orange
COL_ACTIONS   = RGBColor(0x33, 0x66, 0x00)  # green

# Column widths in DXA from the template XML
OBS_WIDTHS   = [675, 3856, 6096]    # Item No., Observations, Comments/Actions
PHOTO_WIDTHS = [5265, 5266]         # left photo, right photo (2-col grid)


# ===========================================================================
# Public API
# ===========================================================================

def generate_report(
    photos: List[PhotoData],
    output_path: str,
    template_path: Optional[str] = None,
    site_name: str = "",
    project_number: str = "",
    inspector_name: str = "",
    site_address: str = "",
) -> str:
    inspection_date = _earliest_date(photos) or date.today()
    report_date     = date.today()
    weather_summary = _weather_summary(photos)
    works_summary   = _works_summary(photos)

    replacements = {
        "SITE_NAME":          site_name or "—",
        "PROJECT_NUMBER":     project_number or "—",
        "INSPECTOR_NAME":     inspector_name or "—",
        "INSPECTION_DATE":    inspection_date.strftime("%d %B %Y"),
        "REPORT_DATE":        report_date.strftime("%d %B %Y"),
        "SITE_ADDRESS":       site_address or "—",
        "TOTAL_PHOTOS":       str(len(photos)),
        "INSPECTION_WEATHER": weather_summary,
        "WORKS_SUMMARY":      works_summary,
    }

    if template_path and os.path.isfile(template_path):
        doc = _load_template(template_path)
        _replace_placeholders(doc, replacements)
        _insert_works_summary(doc, works_summary)
        _populate_observations_table(doc, photos)
        _replace_photo_section(doc, photos)
    else:
        doc = _build_default_doc(photos, replacements, works_summary)

    doc.save(output_path)
    return output_path


# ===========================================================================
# Weather and AI summary
# ===========================================================================

def _weather_summary(photos: List[PhotoData]) -> str:
    summaries = [p.weather.summary() for p in photos if p.weather]
    if not summaries:
        return "Not available"
    return Counter(summaries).most_common(1)[0][0]


def _works_summary(photos: List[PhotoData]) -> str:
    texts = [p.what_inspected.strip() for p in photos if p.what_inspected.strip()]
    if not texts:
        return ""

    combined = " | ".join(texts)

    if _ANTHROPIC_AVAILABLE:
        key = os.environ.get("ANTHROPIC_API_KEY", "")
        if key:
            try:
                client = _anthropic.Anthropic(api_key=key)
                resp = client.messages.create(
                    model="claude-haiku-4-5-20251001",
                    max_tokens=200,
                    messages=[{
                        "role": "user",
                        "content": (
                            "You are drafting a civil site inspection report. "
                            "Based on the following per-photo inspection notes, write one paragraph "
                            "(maximum 100 words) that holistically summarises what was inspected "
                            "during this site visit. Use professional engineering language, "
                            "third-person, past tense. Output only the summary paragraph.\n\n"
                            f"Notes: {combined}"
                        ),
                    }],
                )
                return resp.content[0].text.strip()
            except Exception:
                pass

    # Fallback: truncate combined text to ~100 words
    words = combined.split()
    return " ".join(words[:100]) + ("…" if len(words) > 100 else "")


# ===========================================================================
# Template loader (.docx or .dotx)
# ===========================================================================

def _load_template(path: str) -> Document:
    """Load a .docx or .dotx template.

    .dotx files use content-type 'template.main+xml' which python-docx rejects.
    We rewrite the ZIP in memory, patching the content type to the normal docx
    value, then load the result.
    """
    with zipfile.ZipFile(path, "r") as z:
        ct = z.read("[Content_Types].xml")

    if b"template.main+xml" not in ct:
        return Document(path)

    # Patch: replace template content-type with document content-type
    patched_ct = ct.replace(
        b"wordprocessingml.template.main+xml",
        b"wordprocessingml.document.main+xml",
    )

    tmp = tempfile.mktemp(suffix=".docx")
    try:
        with zipfile.ZipFile(path, "r") as z_in, \
             zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as z_out:
            for item in z_in.infolist():
                data = z_in.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = patched_ct
                z_out.writestr(item, data)
        return Document(tmp)
    finally:
        if os.path.exists(tmp):
            os.unlink(tmp)


# ===========================================================================
# Placeholder replacement
# ===========================================================================

def _replace_placeholders(doc: Document, replacements: dict) -> None:
    def _fix_para(para):
        full = "".join(r.text for r in para.runs)
        new  = PLACEHOLDER_RE.sub(
            lambda m: replacements.get(m.group(1), m.group(0)), full
        )
        if new != full:
            if not para.runs:
                para.add_run(new)
            else:
                for run in para.runs:
                    run.text = ""
                para.runs[0].text = new

    for para in doc.paragraphs:
        _fix_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _fix_para(para)

    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf:
                for para in hf.paragraphs:
                    _fix_para(para)


# ===========================================================================
# Works summary insertion
# ===========================================================================

def _insert_works_summary(doc: Document, summary: str) -> None:
    """Insert the AI works summary as a new paragraph after the
    'proposed works' sentence already in the template body."""
    if not summary:
        return

    body = doc.element.body
    target = None
    for elem in list(body):
        if elem.tag == qn("w:p"):
            text = "".join((t.text or "") for t in elem.iter(qn("w:t")))
            if "proposed works" in text.lower():
                target = elem
                break

    if target is None:
        return

    # Build a plain (non-highlighted) paragraph with the summary
    new_p  = OxmlElement("w:p")
    pPr    = OxmlElement("w:pPr")
    _append_rpr_fonts(pPr, child=True)   # pPr-level rPr for paragraph mark
    new_p.append(pPr)

    r   = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    _set_run_props(rPr)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = summary
    r.append(t)
    new_p.append(r)

    target.addnext(new_p)


# ===========================================================================
# Site Observations table
# ===========================================================================

def _populate_observations_table(doc: Document, photos: List[PhotoData]) -> None:
    table = _find_obs_table(doc)
    if table is None:
        return

    tbl = table._tbl
    # Strip all data rows, keep only the header row (index 0)
    for tr in list(tbl.findall(qn("w:tr")))[1:]:
        tbl.remove(tr)

    item_num = 0
    for idx, photo in enumerate(photos, 1):
        issues = photo.issues_found.strip()
        if not issues:
            continue
        item_num += 1
        what   = photo.what_inspected.strip()
        prefix = f"Photo {idx}"
        if what:
            prefix += f" – {what}"
        obs     = f"{prefix}: {issues}"
        actions = photo.actions_required.strip() or "No contractor action specified."
        _append_obs_row(tbl, str(item_num), obs, actions)

    if item_num == 0:
        _append_obs_row(tbl, "1", "No issues were noted during this inspection.", "—")


def _find_obs_table(doc: Document):
    for table in doc.tables:
        if not table.rows:
            continue
        header_text = " ".join(c.text for c in table.rows[0].cells).lower()
        if "observations" in header_text and "comments" in header_text:
            return table
    return None


def _append_obs_row(tbl_elem, item_no: str, obs: str, actions: str) -> None:
    widths  = OBS_WIDTHS
    texts   = [item_no, obs, actions]
    centers = [True, False, False]

    tr   = OxmlElement("w:tr")
    trPr = OxmlElement("w:trPr")
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), "519")
    trPr.append(trH)
    tr.append(trPr)

    for width, text, center in zip(widths, texts, centers):
        tc   = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr")
        tcW  = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(width))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)
        vA = OxmlElement("w:vAlign")
        vA.set(qn("w:val"), "center")
        tcPr.append(vA)
        tc.append(tcPr)

        p   = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        spc = OxmlElement("w:spacing")
        spc.set(qn("w:before"), "60")
        spc.set(qn("w:after"), "60")
        pPr.append(spc)
        if center:
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "center")
            pPr.append(jc)
        p.append(pPr)

        r   = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        _set_run_props(rPr)
        r.append(rPr)
        t = OxmlElement("w:t")
        if text and (text[0] == " " or text[-1] == " "):
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        r.append(t)
        p.append(r)
        tc.append(p)
        tr.append(tc)

    tbl_elem.append(tr)


# ===========================================================================
# Photo section replacement
# ===========================================================================

def _replace_photo_section(doc: Document, photos: List[PhotoData]) -> None:
    body     = doc.element.body
    children = list(body)

    # Find "Site Photographs" paragraph
    site_ph_idx = None
    for i, elem in enumerate(children):
        if elem.tag == qn("w:p"):
            text = "".join((t.text or "") for t in elem.iter(qn("w:t")))
            if "Site Photographs" in text:
                site_ph_idx = i
                break

    if site_ph_idx is None:
        _append_photo_pages(doc, photos)
        return

    # Preserve the final sectPr (must be last body element for page layout)
    final_sectPr = children[-1] if children[-1].tag == qn("w:sectPr") else None

    for elem in children[site_ph_idx + 1:]:
        if elem is not final_sectPr:
            body.remove(elem)

    _append_photo_pages(doc, photos)


def _append_photo_pages(doc: Document, photos: List[PhotoData]) -> None:
    """Append pages of 4 photos each (2×2 grid): photo row + caption row per pair."""
    if not photos:
        doc.add_paragraph()
        return

    indexed    = list(enumerate(photos, 1))
    page_groups = [indexed[i:i+4] for i in range(0, len(indexed), 4)]

    for page_idx, group in enumerate(page_groups):
        # Spacer paragraph (page break for pages 2+)
        sep = doc.add_paragraph()
        if page_idx > 0:
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            sep.add_run()._r.append(br)

        # Pair photos left/right within this page
        pairs = [
            (group[i], group[i + 1] if i + 1 < len(group) else None)
            for i in range(0, len(group), 2)
        ]

        table = doc.add_table(rows=0, cols=2)
        _configure_photo_table(table)

        for left_entry, right_entry in pairs:
            left_num,  left_photo  = left_entry
            right_num, right_photo = right_entry if right_entry else (None, None)

            # ── Photo row ──────────────────────────────────────────────
            photo_row = table.add_row()
            _set_row_height(photo_row, 4535)          # ≈ 3.15 inches

            lp = photo_row.cells[0]
            rp = photo_row.cells[1]
            _set_cell_width_dxa(lp, PHOTO_WIDTHS[0])
            _set_cell_width_dxa(rp, PHOTO_WIDTHS[1])
            _set_cell_valign(lp, "center")
            _set_cell_valign(rp, "center")

            _put_photo(lp, left_photo, left_num)
            if right_photo:
                _put_photo(rp, right_photo, right_num)

            # ── Caption row ────────────────────────────────────────────
            desc_row = table.add_row()
            ld = desc_row.cells[0]
            rd = desc_row.cells[1]
            _set_cell_width_dxa(ld, PHOTO_WIDTHS[0])
            _set_cell_width_dxa(rd, PHOTO_WIDTHS[1])

            _fill_caption(ld, left_photo, left_num)
            if right_photo:
                _fill_caption(rd, right_photo, right_num)


def _configure_photo_table(table) -> None:
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Table style
    if tblPr.find(qn("w:tblStyle")) is None:
        ts = OxmlElement("w:tblStyle")
        ts.set(qn("w:val"), "TableGrid")
        tblPr.insert(0, ts)

    # Table width
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(PHOTO_WIDTHS)))
    tblW.set(qn("w:type"), "dxa")

    # Column grid
    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for w in PHOTO_WIDTHS:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    tbl_list    = list(tbl)
    tblPr_idx   = tbl_list.index(tblPr)
    tbl.insert(tblPr_idx + 1, grid)


def _put_photo(cell, photo: PhotoData, num: int) -> None:
    """Insert a photo image centred in a table cell."""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        img = PILImage.open(photo.file_path)
        img = ImageOps.exif_transpose(img)
        w, h   = img.size
        ratio  = min(PHOTO_MAX_W_IN * 96 / w, PHOTO_MAX_H_IN * 96 / h, 1.0)
        img    = img.resize(
            (max(1, int(w * ratio)), max(1, int(h * ratio))),
            PILImage.LANCZOS,
        )
        buf = io.BytesIO()
        img.save(buf, "JPEG", quality=85)
        buf.seek(0)
        para.add_run().add_picture(buf, width=Inches(PHOTO_MAX_W_IN))

        if photo.is_duplicate:
            dp = cell.add_paragraph(f"⚠ DUPLICATE of {photo.similar_to}")
            dp.runs[0].font.color.rgb = RGBColor(0xCC, 0, 0)
            dp.runs[0].font.bold      = True
            dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif photo.similar_to:
            sp = cell.add_paragraph(f"≈ Similar to {photo.similar_to}")
            sp.runs[0].font.color.rgb = RGBColor(0xFF, 0x88, 0)
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as exc:
        para.text = f"[Image unavailable: {exc}]"


def _fill_caption(cell, photo: PhotoData, num: int) -> None:
    """Fill a caption cell with photo number, metadata, and inspection notes."""
    PT8 = Pt(8)

    # Photo number header — use the cell's default first paragraph
    first = cell.paragraphs[0]
    first.paragraph_format.space_after = Pt(3)
    rn = first.add_run(f"Photo {num}")
    rn.bold = True
    rn.font.size = PT8
    rn.font.color.rgb = COL_NAV

    def label_value(label: str, value: str) -> None:
        p  = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rl = p.add_run(label)
        rl.bold = True
        rl.font.size = PT8
        rv = p.add_run(value)
        rv.font.size = PT8

    label_value("Date / Time:  ", photo.datetime_label)
    label_value("Coordinates: ", photo.coords_label)
    if photo.direction_degrees is not None:
        label_value("Direction: ", photo.direction_label)
    if photo.weather:
        label_value("Weather: ", photo.weather.summary())

    def inspection_section(heading: str, text: str, color: RGBColor) -> None:
        h  = cell.add_paragraph()
        h.paragraph_format.space_after = Pt(1)
        rh = h.add_run(heading)
        rh.bold = True
        rh.font.size = PT8
        rh.font.color.rgb = color

        b  = cell.add_paragraph()
        b.paragraph_format.space_after = Pt(4)
        rb = b.add_run(text.strip() or "(none noted)")
        rb.font.size = PT8

    inspection_section("What Was Inspected:", photo.what_inspected, COL_INSPECTED)
    inspection_section("Issues Found:",       photo.issues_found,   COL_ISSUES)
    inspection_section("Actions Required:",   photo.actions_required, COL_ACTIONS)


# ===========================================================================
# XML helpers
# ===========================================================================

def _set_run_props(rPr) -> None:
    """Apply standard 9pt Arial to a w:rPr element."""
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:cs"), "Arial")
    rPr.append(rf)
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), "18")   # 18 half-points = 9pt
        rPr.append(el)


def _append_rpr_fonts(pPr, child: bool = False) -> None:
    """Add a standard rPr block to a pPr element (for paragraph-mark formatting)."""
    rPr = OxmlElement("w:rPr")
    _set_run_props(rPr)
    pPr.append(rPr)


def _set_row_height(row, height_dxa: int) -> None:
    tr   = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    trH = trPr.find(qn("w:trHeight"))
    if trH is None:
        trH = OxmlElement("w:trHeight")
        trPr.append(trH)
    trH.set(qn("w:val"), str(height_dxa))
    trH.set(qn("w:hRule"), "exact")


def _set_cell_width_dxa(cell, width_dxa: int) -> None:
    tc   = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")


def _set_cell_valign(cell, val: str) -> None:
    tc   = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    va = tcPr.find(qn("w:vAlign"))
    if va is None:
        va = OxmlElement("w:vAlign")
        tcPr.append(va)
    va.set(qn("w:val"), val)


# ===========================================================================
# Standalone default document (no template)
# ===========================================================================

def _build_default_doc(
    photos: List[PhotoData],
    replacements: dict,
    works_summary: str,
) -> Document:
    doc = Document()

    title = doc.add_heading("SITE INSPECTION REPORT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Header info table
    info = doc.add_table(rows=8, cols=2)
    info.style = "Table Grid"
    for row, (label, val) in zip(info.rows, [
        ("Site / Project:",  replacements["SITE_NAME"]),
        ("Project Number:",  replacements["PROJECT_NUMBER"]),
        ("Site Address:",    replacements["SITE_ADDRESS"]),
        ("Inspector:",       replacements["INSPECTOR_NAME"]),
        ("Inspection Date:", replacements["INSPECTION_DATE"]),
        ("Report Date:",     replacements["REPORT_DATE"]),
        ("Total Photos:",    replacements["TOTAL_PHOTOS"]),
        ("Weather:",         replacements["INSPECTION_WEATHER"]),
    ]):
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = val

    doc.add_paragraph()

    # Works summary
    if works_summary:
        h = doc.add_heading("Summary of Inspected Works", level=1)
        p = doc.add_paragraph(works_summary)

    doc.add_paragraph()
    doc.add_heading("Site Observations", level=1)

    # Observations table
    obs = doc.add_table(rows=1, cols=3)
    obs.style = "Table Grid"
    for cell, hdr in zip(obs.rows[0].cells,
                         ["Item No.", "Observations", "Comments / Actions"]):
        cell.text = hdr
        cell.paragraphs[0].runs[0].bold = True

    item_num = 0
    for idx, photo in enumerate(photos, 1):
        issues = photo.issues_found.strip()
        if not issues:
            continue
        item_num += 1
        what    = photo.what_inspected.strip()
        obs_txt = f"Photo {idx}" + (f" – {what}" if what else "") + f": {issues}"
        actions = photo.actions_required.strip() or "No contractor action specified."
        row = obs.add_row()
        row.cells[0].text = str(item_num)
        row.cells[1].text = obs_txt
        row.cells[2].text = actions

    if item_num == 0:
        row = obs.add_row()
        row.cells[0].text = "1"
        row.cells[1].text = "No issues were noted during this inspection."
        row.cells[2].text = "—"

    doc.add_paragraph()
    doc.add_paragraph(
        "Please contact the undersigned if you require clarification of any of the items above."
    )
    doc.add_paragraph()
    doc.add_paragraph("Regards,")
    doc.add_paragraph()
    pn = doc.add_paragraph()
    r  = pn.add_run(replacements["INSPECTOR_NAME"])
    r.bold = True
    doc.add_paragraph("On behalf of Meinhardt Australia")

    # ── Photo appendix ──────────────────────────────────────────────────────
    pb = doc.add_paragraph()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    pb.add_run()._r.append(br)

    h2 = doc.add_heading("Site Photographs", level=1)

    _append_photo_pages(doc, photos)

    return doc


# ===========================================================================
# Date utility
# ===========================================================================

def _earliest_date(photos: List[PhotoData]) -> Optional[date]:
    dates = [p.datetime_taken.date() for p in photos if p.datetime_taken]
    return min(dates) if dates else None
